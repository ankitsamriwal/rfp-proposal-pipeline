"""Pipeline stages. Each stage builds a prompt, calls the provider, writes artifacts."""
import csv
import os
import re

SYSTEM = {
    "exec_summary": "TASK: exec_summary\nYou are a presales analyst. Write a concise executive summary of this RFP in markdown.",
    "gap_analysis": "TASK: gap_analysis\nYou are a presales analyst. Produce a gap analysis between this RFP and a standard Dynamics 365 Business Central delivery template, as a markdown table.",
    "clarification_log": "TASK: clarification_log\nYou are a presales analyst. Draft a clarification-question log (markdown table: ID, question, status) for the ambiguities in this RFP.",
    "risk_flags": "TASK: risk_flags\nYou are a contracts-aware presales analyst. Flag contractual risks (bank guarantees, performance bonds, insurance, penalties/liquidated damages, arbitration, free licences, warranties, indemnities) as a markdown table: risk, severity, why it matters.",
    "compliance_matrix": "TASK: compliance_matrix\nYou are a bid manager. Extract every mandatory requirement ('shall'/'must'/'required') into a compliance matrix markdown table: ID, requirement, response, owner.",
    "proposal_skeleton": "TASK: proposal_skeleton\nYou are a proposal manager. Output a numbered proposal skeleton for responding to this RFP.",
}


def _run_stage(provider, task, rfp_text):
    prompt = f"Analyse the following RFP document.\n\nRFP TEXT:\n{rfp_text}"
    return provider.complete(SYSTEM[task], prompt)


def _write(run_dir, name, content):
    path = os.path.join(run_dir, name)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def stage_analyst(provider, rfp_text, run_dir):
    """Stage 1: exec summary, gap analysis, clarification log, risk flags."""
    outputs = []
    outputs.append(_write(run_dir, "01_executive_summary.md", _run_stage(provider, "exec_summary", rfp_text)))
    outputs.append(_write(run_dir, "02_gap_analysis.md", _run_stage(provider, "gap_analysis", rfp_text)))
    outputs.append(_write(run_dir, "03_clarification_log.md", _run_stage(provider, "clarification_log", rfp_text)))
    outputs.append(_write(run_dir, "04_risk_flags.md", _run_stage(provider, "risk_flags", rfp_text)))
    return outputs


def _md_table_to_rows(md):
    rows = []
    for line in md.splitlines():
        line = line.strip()
        if line.startswith("|") and not re.match(r"^\|[\s\-|]+\|$", line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
    return rows


def stage_compliance(provider, rfp_text, run_dir, out_name="05_compliance_matrix"):
    """Stage 2: compliance matrix as CSV + Markdown."""
    md = _run_stage(provider, "compliance_matrix", rfp_text)
    md_path = _write(run_dir, f"{out_name}.md", md)
    rows = _md_table_to_rows(md)
    csv_path = os.path.join(run_dir, f"{out_name}.csv")
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if rows:
            writer.writerow(rows[0])
            writer.writerows(rows[1:])
        else:
            writer.writerow(["ID", "Requirement", "Response", "Owner"])
    return [md_path, csv_path]


def stage_proposal(provider, rfp_text, run_dir, out_name="06_proposal_skeleton"):
    """Stage 3: proposal skeleton as DOCX."""
    import docx
    skeleton_md = _run_stage(provider, "proposal_skeleton", rfp_text)
    md_path = _write(run_dir, f"{out_name}.md", skeleton_md)

    doc = docx.Document()
    doc.add_heading("Proposal Skeleton", level=0)
    for line in skeleton_md.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            continue  # already added the document heading
        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2), level=1)
            doc.add_paragraph("[To be drafted by the proposal team]")
        elif line.startswith("**") and line.endswith("**"):
            doc.add_paragraph(line.strip("*"))
        else:
            doc.add_paragraph(line)
    docx_path = os.path.join(run_dir, f"{out_name}.docx")
    doc.save(docx_path)
    return [md_path, docx_path]
