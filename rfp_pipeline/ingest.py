"""Ingest RFP documents (PDF, DOCX, TXT, MD) from a folder and extract text."""
import os


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".txt", ".md"):
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    raise ValueError(f"Unsupported file type: {ext} ({path})")


def _extract_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(path):
    import docx
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def ingest_folder(folder):
    """Return {filename: text} for every supported document in the folder."""
    docs = {}
    for name in sorted(os.listdir(folder)):
        path = os.path.join(folder, name)
        if not os.path.isfile(path):
            continue
        if os.path.splitext(name)[1].lower() in (".pdf", ".docx", ".txt", ".md"):
            docs[name] = extract_text(path)
    if not docs:
        raise FileNotFoundError(f"No RFP documents (.pdf/.docx/.txt/.md) found in {folder}")
    return docs
