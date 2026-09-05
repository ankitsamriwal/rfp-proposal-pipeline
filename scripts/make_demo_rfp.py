"""Generate the synthetic demo RFP: Gulf Crescent Trading LLC (D365 BC implementation).

The company is fictional. The document deliberately includes the contractual
hooks the pipeline is built to flag: bank guarantee, insurance, penalties,
arbitration, Arabic requirements, and a request for free licences.
"""
import os

import docx

OUT = os.path.join(os.path.dirname(__file__), "..", "demo_rfp", "gulf_crescent_trading_llc_rfp.docx")

SECTIONS = [
    ("Request for Proposal - ERP Implementation (Microsoft Dynamics 365 Business Central)", None),
    ("Issued by: Gulf Crescent Trading LLC", None),
    ("1. Introduction", [
        "Gulf Crescent Trading LLC ('the Company'), a trading and distribution business headquartered in Dubai, UAE, invites qualified implementation partners to submit proposals for the implementation of Microsoft Dynamics 365 Business Central.",
        "The Company operates three warehouses in Jebel Ali and imports consumer electronics for distribution across the GCC.",
    ]),
    ("2. Scope of Work", [
        "The selected partner shall implement Dynamics 365 Business Central covering finance, procurement, inventory management, and sales order processing.",
        "The partner must migrate opening balances and two years of transactional history from the legacy Tally system.",
        "The solution shall integrate with the Company's existing e-commerce platform and its 3PL warehouse management system.",
        "The partner is required to deliver end-user training for approximately 40 users, including training material in both English and Arabic.",
        "The user interface and all customer-facing reports must be available in Arabic as well as English.",
    ]),
    ("3. Timeline", [
        "The implementation shall go live no later than 1 March 2027.",
        "Proposals must be submitted by 15 October 2026 at 17:00 GST.",
    ]),
    ("4. Commercial and Contractual Requirements", [
        "The bidder shall provide a bank guarantee equal to 10% of the total contract value, valid until final acceptance, issued by a UAE-licensed bank.",
        "A performance bond of 5% of contract value is required upon contract signing.",
        "The bidder must maintain professional indemnity insurance of not less than AED 2,000,000 and third-party liability insurance for the duration of the engagement.",
        "Delay in delivery beyond the agreed go-live date shall attract liquidated damages of 0.5% of contract value per week, up to a maximum of 10%.",
        "Any dispute arising from the contract shall be settled by arbitration under DIAC rules, seated in Dubai, conducted in English.",
        "The bidder shall provide all end-user licences free of charge for the first year of operation.",
        "The bidder shall provide a warranty period of 12 months following go-live, including bug fixes at no additional cost.",
        "The bidder shall indemnify the Company against any claims arising from intellectual property infringement by the delivered solution.",
    ]),
    ("5. Proposal Requirements", [
        "The bidder shall submit a technical proposal describing methodology, solution architecture, team composition, and relevant case studies.",
        "The bidder must submit a separate commercial proposal with a fixed-price implementation quote and itemised licensing costs.",
        "The bidder is required to complete the compliance matrix in Appendix A, responding to every mandatory requirement.",
        "Bidders must have at least three referenceable Dynamics 365 Business Central implementations in the UAE or wider GCC.",
    ]),
    ("6. Evaluation", [
        "Proposals will be evaluated 60% on technical merit and 40% on commercial competitiveness.",
        "The Company is not obliged to accept the lowest-priced or any proposal.",
    ]),
]


def main():
    doc = docx.Document()
    for heading, paragraphs in SECTIONS:
        if paragraphs is None:
            doc.add_heading(heading, level=0)
        else:
            doc.add_heading(heading, level=1)
            for p in paragraphs:
                doc.add_paragraph(p)
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
